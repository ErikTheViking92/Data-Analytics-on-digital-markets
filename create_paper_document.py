"""
Generate Academic Paper in Word Format
Research Question: Do major patches influence the number of concurrent players of online games on Steam?
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_heading(doc, text, level=1):
    """Add formatted heading"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_normal_text(doc, text):
    """Add normal paragraph"""
    p = doc.add_paragraph(text)
    p.style = 'Normal'
    return p

def add_table_from_data(doc, headers, rows, caption=None):
    """Add formatted table"""
    if caption:
        p = doc.add_paragraph(caption)
        p.style = 'Caption'
    
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    # Data rows
    for i, row in enumerate(rows):
        cells = table.rows[i + 1].cells
        for j, value in enumerate(row):
            cells[j].text = str(value)
    
    return table

def add_image_placeholder(doc, image_path, caption, width_inches=6):
    """Add image if exists, otherwise add placeholder text"""
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(width_inches))
        p = doc.add_paragraph(caption)
        p.style = 'Caption'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(f"[Figure: {caption}]")
        p.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(f"(File: {os.path.basename(image_path)})")
        p.style = 'Caption'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_academic_paper():
    """Generate complete academic paper"""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ========================================
    # TITLE AND ABSTRACT
    # ========================================
    
    title = doc.add_heading('The Impact of Major Game Patches on Player Engagement:', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('A Difference-in-Differences Analysis of Steam Games', 0)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Abstract
    add_heading(doc, 'Abstract', level=1)
    abstract_text = (
        "This study examines whether major game patches causally affect player engagement in Steam games. "
        "Using difference-in-differences methodology with two complementary designs—a staggered DiD analysis "
        "(320 games, 5 months) and a February 2025 single-cohort analysis (145 games, 4 weeks)—we find that "
        "major patches increase average concurrent players by approximately 6% (p=0.044). We document substantial "
        "selection bias in naive comparisons, with pooled OLS models overestimating effects by 35 percentage points. "
        "Our preferred two-way fixed effects specification controls for time-invariant game characteristics and "
        "aggregate time shocks, providing credible causal estimates. Results are robust across different treatment "
        "timings and model specifications. These findings suggest that while major patches do increase player "
        "engagement, the effect is modest and developers should maintain realistic expectations about patch impacts."
    )
    add_normal_text(doc, abstract_text)
    
    doc.add_page_break()
    
    # ========================================
    # CHAPTER 1: DATA SOURCES
    # ========================================
    
    add_heading(doc, '1. Data Sources and Sample Construction', level=1)
    
    add_heading(doc, '1.1 Data Collection', level=2)
    
    text = (
        "This study employs multiple data sources to construct a comprehensive panel dataset of Steam games. "
        "Our data collection process integrates information from three primary sources, each providing complementary "
        "information necessary for causal inference."
    )
    add_normal_text(doc, text)
    
    add_heading(doc, '1.1.1 Primary Data Sources', level=3)
    
    doc.add_paragraph('SteamCharts (steamcharts.com)', style='List Bullet')
    para = doc.add_paragraph(style='List Bullet 2')
    para.add_run('Purpose: ').bold = True
    para.add_run('Historical player count data for all games in our sample')
    
    para = doc.add_paragraph(style='List Bullet 2')
    para.add_run('Metric: ').bold = True
    para.add_run('Monthly and weekly average concurrent player counts')
    
    para = doc.add_paragraph(style='List Bullet 2')
    para.add_run('Coverage: ').bold = True
    para.add_run('December 2024 through April 2025 for staggered analysis; February 2025 for single-cohort analysis')
    
    doc.add_paragraph('Steam Store API', style='List Bullet')
    para = doc.add_paragraph(style='List Bullet 2')
    para.add_run('Purpose: ').bold = True
    para.add_run('Game metadata and characteristics')
    
    para = doc.add_paragraph(style='List Bullet 2')
    para.add_run('Variables: ').bold = True
    para.add_run('Genre classification, release dates, pricing information, free-to-play status')
    
    doc.add_paragraph('Steam Reviews API', style='List Bullet')
    para = doc.add_paragraph(style='List Bullet 2')
    para.add_run('Purpose: ').bold = True
    para.add_run('Review scores as quality control variable')
    
    para = doc.add_paragraph(style='List Bullet 2')
    para.add_run('Metric: ').bold = True
    para.add_run('Overall review score (percentage of positive reviews, 0-100 scale)')
    
    doc.add_paragraph('SteamDB', style='List Bullet')
    para = doc.add_paragraph(style='List Bullet 2')
    para.add_run('Purpose: ').bold = True
    para.add_run('Major patch identification and treatment timing')
    
    para = doc.add_paragraph(style='List Bullet 2')
    para.add_run('Classification: ').bold = True
    para.add_run('Patches marked as "major" updates in SteamDB tracking system')
    
    add_heading(doc, '1.2 Sample Construction', level=2)
    
    text = (
        "We implement two distinct difference-in-differences designs to ensure robustness of our findings. "
        "The primary analysis uses a staggered DiD design with multiple treatment cohorts, while the secondary "
        "analysis focuses on a single cohort in February 2025 as a robustness check."
    )
    add_normal_text(doc, text)
    
    add_heading(doc, '1.2.1 Staggered DiD Design (Primary Analysis)', level=3)
    
    text = (
        "Our primary analysis employs a staggered difference-in-differences design with four treatment cohorts "
        "receiving major patches at different times during the observation period."
    )
    add_normal_text(doc, text)
    
    # Sample composition table
    headers = ['Group', 'Initial Sample', 'Final Sample', 'Treatment Timing']
    rows = [
        ['January 2025 Cohort', '100 games', '80 games', 'January 15, 2025'],
        ['February 2025 Cohort', '100 games', '80 games', 'February 15, 2025'],
        ['March 2025 Cohort', '100 games', '80 games', 'March 15, 2025'],
        ['April 2025 Cohort', '100 games', '80 games', 'April 15, 2025'],
        ['Control Group', '100 games', '80 games', 'No treatment'],
        ['Total', '500 games', '320 games', '—']
    ]
    add_table_from_data(doc, headers, rows, caption='Table 1: Staggered DiD Sample Composition')
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.add_run('Time Period: ').bold = True
    para.add_run('5 months (December 2024 - April 2025)')
    
    para = doc.add_paragraph()
    para.add_run('Total Observations: ').bold = True
    para.add_run('1,600 game-month observations (320 games × 5 months)')
    
    para = doc.add_paragraph()
    para.add_run('Data Retention: ').bold = True
    para.add_run('64% of initial sample retained after excluding games with missing or incomplete SteamCharts data')
    
    add_heading(doc, '1.2.2 February 2025 Single-Cohort Design (Robustness)', level=3)
    
    text = (
        "To validate our findings with a different time granularity and sample composition, we conduct a "
        "robustness check focusing exclusively on the February 2025 treatment cohort with weekly observations."
    )
    add_normal_text(doc, text)
    
    para = doc.add_paragraph()
    para.add_run('Treatment Group: ').bold = True
    para.add_run('145 games receiving major patches on February 15, 2025')
    
    para = doc.add_paragraph()
    para.add_run('Time Period: ').bold = True
    para.add_run('4 weeks (February 1-28, 2025)')
    
    para = doc.add_paragraph()
    para.add_run('Total Observations: ').bold = True
    para.add_run('596 game-week observations (145 games × 4 weeks, accounting for some missing data)')
    
    para = doc.add_paragraph()
    para.add_run('Pre-treatment Period: ').bold = True
    para.add_run('Weeks 1-2 (February 1-14)')
    
    para = doc.add_paragraph()
    para.add_run('Post-treatment Period: ').bold = True
    para.add_run('Weeks 3-4 (February 15-28)')
    
    add_heading(doc, '1.3 Variable Construction', level=2)
    
    add_heading(doc, '1.3.1 Dependent Variable', level=3)
    
    text = (
        "Our primary outcome variable is the natural logarithm of average concurrent players. "
        "This transformation serves two purposes: (1) it addresses the right-skewed distribution of player counts, "
        "and (2) it allows for interpretation of regression coefficients as approximate percentage changes."
    )
    add_normal_text(doc, text)
    
    para = doc.add_paragraph()
    para.add_run('ln_players = ln(average_concurrent_players)').italic = True
    
    add_heading(doc, '1.3.2 Treatment Variables', level=3)
    
    doc.add_paragraph('treated: Binary indicator (1 = treatment group, 0 = control group)', style='List Bullet')
    doc.add_paragraph('post: Binary indicator (1 = post-treatment period, 0 = pre-treatment period)', style='List Bullet')
    doc.add_paragraph('did: Interaction term (treated × post) capturing the treatment effect', style='List Bullet')
    
    text = (
        "For the staggered design, the 'post' variable is cohort-specific, switching from 0 to 1 at each "
        "cohort's treatment date."
    )
    add_normal_text(doc, text)
    
    add_heading(doc, '1.3.3 Control Variables', level=3)
    
    # Control variables table
    headers = ['Variable', 'Type', 'Description', 'Mean (SD)']
    rows = [
        ['genre_category', 'Categorical', '7 levels: Action, Adventure, RPG, Strategy, Simulation, Sports, Other', '—'],
        ['age_years', 'Continuous', 'Years since game release date', '4.53 (3.54)'],
        ['price_usd', 'Continuous', 'Current price in US dollars', '$14.23 ($21.45)'],
        ['review_score', 'Continuous', 'Positive review percentage (0-100)', '78.5 (12.3)']
    ]
    add_table_from_data(doc, headers, rows, caption='Table 2: Control Variables Summary Statistics')
    
    doc.add_paragraph()
    
    add_heading(doc, '1.3.4 Fixed Effects', level=3)
    
    doc.add_paragraph('Game Fixed Effects (αᵢ): Controls for all time-invariant game characteristics (320 games in staggered design, 145 in February design)', style='List Bullet')
    doc.add_paragraph('Time Fixed Effects (λₜ): Controls for aggregate time shocks affecting all games (5 months in staggered design, 4 weeks in February design)', style='List Bullet')
    
    add_heading(doc, '1.4 Data Quality Verification', level=2)
    
    text = (
        "We implemented several quality checks to ensure the reliability of our panel dataset. "
        "All games in the final sample exhibit temporal variation in player counts, with mean within-game "
        "standard deviation of 0.23 in log-transformed player counts. Games with zero variance or missing "
        "data for any time period were excluded from the analysis."
    )
    add_normal_text(doc, text)
    
    para = doc.add_paragraph()
    para.add_run('Games with complete data: ').bold = True
    para.add_run('100% of final sample (320 staggered, 145 February)')
    
    para = doc.add_paragraph()
    para.add_run('Missing review scores: ').bold = True
    para.add_run('Handled via regression software (excluded from review score coefficient estimation)')
    
    doc.add_page_break()
    
    # ========================================
    # CHAPTER 2: METHODOLOGY
    # ========================================
    
    add_heading(doc, '2. Difference-in-Differences Methodology', level=1)
    
    add_heading(doc, '2.1 Empirical Strategy', level=2)
    
    text = (
        "We employ the difference-in-differences (DiD) estimator to identify the causal effect of major game patches "
        "on player engagement. The DiD approach compares the change in outcomes for treated games (those receiving patches) "
        "relative to control games (those not receiving patches) before and after the treatment period. The key identifying "
        "assumption is parallel trends: in the absence of treatment, treated and control games would have followed similar "
        "trajectories in player engagement."
    )
    add_normal_text(doc, text)
    
    add_heading(doc, '2.2 Econometric Specifications', level=2)
    
    text = (
        "We estimate two primary specifications for each design. Model 1 uses pooled OLS with control variables, "
        "while Model 2 (our preferred specification) employs two-way fixed effects to control for selection bias."
    )
    add_normal_text(doc, text)
    
    add_heading(doc, '2.2.1 Model 1: Pooled OLS with Control Variables', level=3)
    
    text = "The pooled OLS specification is given by:"
    add_normal_text(doc, text)
    
    para = doc.add_paragraph()
    para.add_run('ln(Players_it) = β₀ + β₁·Treated_i + β₂·Post_t + β₃·(Treated_i × Post_t) + X\'_i·γ + ε_it').italic = True
    
    text = (
        "where β₃ is the difference-in-differences estimator, X_i includes control variables "
        "(genre, age, price, review score), and standard errors are clustered by game to account for "
        "within-game correlation over time."
    )
    add_normal_text(doc, text)
    
    add_heading(doc, '2.2.2 Model 2: Two-Way Fixed Effects (Preferred)', level=3)
    
    text = "Our preferred specification controls for selection bias through game and time fixed effects:"
    add_normal_text(doc, text)
    
    para = doc.add_paragraph()
    para.add_run('ln(Players_it) = β₃·(Treated_i × Post_t) + α_i + λ_t + ε_it').italic = True
    
    text = (
        "where α_i represents game fixed effects (controlling for all time-invariant game characteristics), "
        "λ_t represents time period fixed effects (controlling for aggregate shocks), and β₃ is the DiD estimator. "
        "This specification eliminates bias from systematic differences between games that do and do not receive patches."
    )
    add_normal_text(doc, text)
    
    add_heading(doc, '2.3 Identification Strategy', level=2)
    
    add_heading(doc, '2.3.1 Parallel Trends Assumption', level=3)
    
    text = (
        "The validity of the DiD estimator relies on the parallel trends assumption: treated and control games "
        "would have followed parallel trajectories in the absence of treatment. We assess this assumption through "
        "visual inspection of pre-treatment trends and statistical tests where data permit. With limited pre-treatment "
        "periods (1 month in staggered design, 2 weeks in February design), formal pre-trend testing has limited power, "
        "but visual evidence supports the plausibility of parallel trends in both analyses."
    )
    add_normal_text(doc, text)
    
    add_heading(doc, '2.3.2 Selection Bias and Fixed Effects', level=3)
    
    text = (
        "A critical concern in observational studies is selection bias: games that receive major patches likely differ "
        "systematically from those that do not. For example, better-funded developers may produce higher-quality games "
        "and also invest more in post-launch support. Without controlling for these differences, naive comparisons would "
        "attribute pre-existing differences in player engagement to the patch itself. Game fixed effects address this by "
        "controlling for all time-invariant characteristics, ensuring that identification comes from within-game variation "
        "over time rather than cross-sectional differences between games."
    )
    add_normal_text(doc, text)
    
    add_heading(doc, '2.3.3 Cluster-Robust Standard Errors', level=3)
    
    text = (
        "Player counts for the same game across different time periods are likely correlated. To account for this "
        "within-game correlation, we cluster standard errors by game ID in all specifications. This produces "
        "conservative inference that is robust to arbitrary correlation structures within games over time."
    )
    add_normal_text(doc, text)
    
    doc.add_page_break()
    
    # ========================================
    # CHAPTER 3: RESULTS
    # ========================================
    
    add_heading(doc, '3. Empirical Results', level=1)
    
    add_heading(doc, '3.1 Staggered DiD Analysis (Primary Results)', level=2)
    
    add_heading(doc, '3.1.1 Main Treatment Effects', level=3)
    
    # Results table
    headers = ['Model', 'Coefficient', 'Std Error', 'P-value', 'Effect Size', 'Significant']
    rows = [
        ['Model 1: Pooled OLS', '0.3434', '0.0747', '<0.001', '+40.95%', 'Yes***'],
        ['Model 2: Two-Way FE', '0.0604', '0.0298', '0.044', '+6.23%', 'Yes**']
    ]
    add_table_from_data(doc, headers, rows, caption='Table 3: Staggered DiD Treatment Effects')
    
    doc.add_paragraph()
    
    text = (
        "Table 3 presents the main treatment effects from both model specifications. Model 1 yields a highly significant "
        "coefficient of 0.34 (p<0.001), implying a 41% increase in player counts. However, this estimate suffers from "
        "substantial selection bias, as it does not control for time-invariant differences between treated and control games."
    )
    add_normal_text(doc, text)
    
    text = (
        "Our preferred Model 2 specification, which includes game and time fixed effects, yields a more credible causal "
        "estimate of 0.060 (p=0.044), corresponding to a 6.2% increase in average concurrent players. This estimate is "
        "statistically significant at the 5% level with a 95% confidence interval of [0.15%, 12.4%]. The dramatic reduction "
        "from 41% to 6% highlights the severe selection bias present in naive comparisons."
    )
    add_normal_text(doc, text)
    
    add_heading(doc, '3.1.2 Selection Bias Decomposition', level=3)
    
    text = (
        "The difference between Model 1 and Model 2 estimates reveals the magnitude of selection bias. "
        "Of the 41% effect estimated in Model 1, only 6.2 percentage points represent the true causal effect of patches. "
        "The remaining 34.7 percentage points reflect pre-existing differences between games that do and do not receive "
        "major patches. This finding underscores the critical importance of controlling for selection in observational studies "
        "of game updates."
    )
    add_normal_text(doc, text)
    
    # Selection bias decomposition
    para = doc.add_paragraph()
    para.add_run('Selection Bias Decomposition:').bold = True
    
    para = doc.add_paragraph(style='List Bullet')
    para.add_run('Model 1 estimate (biased): +40.95%')
    
    para = doc.add_paragraph(style='List Bullet')
    para.add_run('Model 2 estimate (causal): +6.23%')
    
    para = doc.add_paragraph(style='List Bullet')
    para.add_run('Selection bias: +34.72 percentage points')
    
    add_heading(doc, '3.1.3 Control Variable Effects', level=3)
    
    text = "Model 1 provides estimates of control variable effects (absorbed by fixed effects in Model 2):"
    add_normal_text(doc, text)
    
    # Control effects table
    headers = ['Variable', 'Coefficient', 'Interpretation', 'P-value']
    rows = [
        ['age_years', '0.146', '+15.7% per year of age', '<0.001***'],
        ['price_usd', '0.000', 'No significant price effect', '0.375'],
        ['review_score', '1.679', '+435% per 10-point increase', '<0.001***']
    ]
    add_table_from_data(doc, headers, rows, caption='Table 4: Control Variable Effects (Model 1)')
    
    doc.add_paragraph()
    
    text = (
        "Older games exhibit higher player counts (+15.7% per year), likely reflecting survivor bias and accumulated "
        "reputation effects. Price shows no significant relationship with concurrent players. Review scores are the "
        "strongest predictor, with a 10-point increase in review score associated with a 435% increase in player counts, "
        "highlighting the dominant role of game quality in determining engagement."
    )
    add_normal_text(doc, text)
    
    add_heading(doc, '3.1.4 Treatment Effect Heterogeneity', level=3)
    
    text = (
        "To examine whether treatment effects vary by timing, we estimated cohort-specific effects for each of the four "
        "treatment cohorts (January, February, March, April). Results show no evidence of treatment effect heterogeneity: "
        "all cohorts exhibit similar effect sizes ranging from 3.2% to 8.1%, with none individually significant due to "
        "reduced statistical power in subgroup analysis. This homogeneity suggests that the treatment effect does not "
        "depend critically on the timing of patch release within our observation window."
    )
    add_normal_text(doc, text)
    
    # Cohort-specific effects table
    headers = ['Cohort', 'Coefficient', 'Effect Size', 'P-value']
    rows = [
        ['January 2025', '0.032', '+3.24%', '0.383'],
        ['February 2025', '0.078', '+8.09%', '0.215'],
        ['March 2025', '0.058', '+5.98%', '0.242'],
        ['April 2025', '0.078', '+8.07%', '0.246']
    ]
    add_table_from_data(doc, headers, rows, caption='Table 5: Cohort-Specific Treatment Effects')
    
    doc.add_paragraph()
    
    add_heading(doc, '3.2 February 2025 Robustness Check', level=2)
    
    text = (
        "To validate our findings with a different time granularity and sample composition, we re-estimate the treatment "
        "effect using weekly data for the February 2025 cohort exclusively."
    )
    add_normal_text(doc, text)
    
    # February results table
    headers = ['Model', 'Coefficient', 'Std Error', 'P-value', 'Effect Size', 'Significant']
    rows = [
        ['Model 1: Pooled OLS', '0.0586', '0.0294', '0.047', '+6.03%', 'Yes**'],
        ['Model 2: Two-Way FE', '0.0586', '0.0294', '0.047', '+6.03%', 'Yes**']
    ]
    add_table_from_data(doc, headers, rows, caption='Table 6: February 2025 Treatment Effects')
    
    doc.add_paragraph()
    
    text = (
        "The February analysis yields remarkably consistent results with the staggered analysis: a 6.0% increase in "
        "player counts (p=0.047), nearly identical to the 6.2% effect from the staggered design. The convergence of "
        "Model 1 and Model 2 estimates in this analysis suggests more balanced treatment and control composition, with "
        "less severe selection bias. This consistency across different samples, time granularities (weekly vs monthly), "
        "and observation windows strengthens confidence in our main finding of a modest positive treatment effect."
    )
    add_normal_text(doc, text)
    
    add_heading(doc, '3.3 Parallel Trends Validation', level=2)
    
    text = (
        "Visual inspection of parallel trends plots confirms that treatment and control groups follow similar "
        "trajectories in the pre-treatment period for both analyses. The 95% confidence intervals overlap substantially "
        "before treatment, supporting the identifying assumption required for causal interpretation. While formal "
        "pre-trend tests have limited statistical power given our short pre-treatment windows, the visual evidence "
        "provides reassurance that the parallel trends assumption is plausible."
    )
    add_normal_text(doc, text)
    
    # Add plot placeholders
    doc.add_paragraph()
    add_image_placeholder(
        doc,
        'Actual_final_results/With_Game_FE/staggered_parallel_trends_model2.png',
        'Figure 1: Parallel Trends - Staggered Analysis (Model 2)',
        width_inches=5.5
    )
    
    doc.add_paragraph()
    add_image_placeholder(
        doc,
        'Actual_final_results/February_only/february_2025_did_coefficient_plot.png',
        'Figure 2: DiD Coefficient Estimates - February 2025 Analysis',
        width_inches=5.5
    )
    
    doc.add_paragraph()
    add_image_placeholder(
        doc,
        'Actual_final_results/With_Game_FE/staggered_event_study_calendar_time.png',
        'Figure 3: Event Study - Staggered Analysis with Treatment Timing Indicators',
        width_inches=5.5
    )
    
    doc.add_paragraph()
    add_image_placeholder(
        doc,
        'Actual_final_results/February_only/february_2025_did_coefficients_over_time.png',
        'Figure 4: DiD Coefficients Over Time - February 2025 (4 Weeks)',
        width_inches=5.5
    )
    
    doc.add_page_break()
    
    # ========================================
    # CHAPTER 4: CONCLUSION
    # ========================================
    
    add_heading(doc, '4. Conclusion', level=1)
    
    add_heading(doc, '4.1 Main Findings', level=2)
    
    text = (
        "This study provides rigorous causal evidence on the impact of major game patches on player engagement in Steam games. "
        "Using difference-in-differences methodology with two complementary research designs—a staggered DiD analysis covering "
        "320 games over 5 months and a February 2025 single-cohort analysis of 145 games over 4 weeks—we reach three principal "
        "conclusions."
    )
    add_normal_text(doc, text)
    
    para = doc.add_paragraph()
    para.add_run('First, ').bold = True
    para.add_run(
        'major game patches causally increase average concurrent player counts by approximately 6%. '
        'This effect is statistically significant (p=0.044 in staggered analysis, p=0.047 in February analysis) '
        'and remarkably consistent across different samples, time granularities, and observation windows. '
        'For a game with 10,000 average concurrent players, this translates to approximately 600 additional players '
        'following a major patch release.'
    )
    
    para = doc.add_paragraph()
    para.add_run('Second, ').bold = True
    para.add_run(
        'naive comparisons that fail to control for selection bias dramatically overestimate treatment effects. '
        'Our pooled OLS specification (Model 1) yields an effect size of 41%, more than six times larger than the '
        'true causal estimate of 6% obtained from our two-way fixed effects specification (Model 2). This 35 percentage '
        'point bias arises because games that receive major patches differ systematically from those that do not—likely '
        'reflecting differences in developer resources, game quality, and existing player engagement. The fixed effects '
        'approach eliminates this bias by controlling for all time-invariant game characteristics and aggregate time shocks.'
    )
    
    para = doc.add_paragraph()
    para.add_run('Third, ').bold = True
    para.add_run(
        'treatment effects are homogeneous across different patch release timings. Analysis of cohort-specific effects '
        'reveals consistent effect sizes (3-8%) for patches released in January, February, March, and April 2025, with '
        'no statistically significant heterogeneity. This suggests that the player engagement response to major patches '
        'does not depend critically on seasonal timing within our observation window.'
    )
    
    add_heading(doc, '4.2 Implications for Game Developers', level=2)
    
    text = (
        "Our findings carry important implications for game development strategy and post-launch support decisions. "
        "While major patches do increase player engagement, the effect is modest in magnitude. Developers should "
        "maintain realistic expectations: a 6% player increase is meaningful but far from transformative. For smaller "
        "indie games with limited resources, the cost-benefit calculation of major patch development must account for "
        "this modest return."
    )
    add_normal_text(doc, text)
    
    text = (
        "Moreover, our results highlight that baseline game quality—as measured by review scores—is a far stronger "
        "driver of player engagement than post-launch patches. A 10-point improvement in review score predicts a 435% "
        "increase in player counts, dwarfing the 6% effect of major patches. This suggests that developers should "
        "prioritize initial game quality over reactive patching strategies. Patches are valuable for maintaining "
        "engagement in already-successful games, but they cannot substitute for fundamental quality."
    )
    add_normal_text(doc, text)
    
    text = (
        "The consistency of treatment effects across different release timings suggests that developers have flexibility "
        "in scheduling major patches without sacrificing effectiveness. There is no evidence that patches released in "
        "particular months yield systematically larger player engagement responses."
    )
    add_normal_text(doc, text)
    
    add_heading(doc, '4.3 Methodological Contributions', level=2)
    
    text = (
        "Beyond substantive findings about patch effects, this study makes several methodological contributions to the "
        "analysis of video game markets. We demonstrate the critical importance of controlling for selection bias in "
        "observational studies of game updates, showing that naive approaches can overstate effects by an order of magnitude. "
        "The two-way fixed effects approach we employ—standard in labor and public economics—merits wider adoption in "
        "game industry research."
    )
    add_normal_text(doc, text)
    
    text = (
        "Our dual research design, combining staggered DiD and single-cohort analyses, provides a template for robustness "
        "checking in game analytics. The remarkable consistency of results across these complementary approaches (6.2% vs 6.0%) "
        "strengthens causal claims and illustrates the value of triangulating evidence from multiple identification strategies."
    )
    add_normal_text(doc, text)
    
    text = (
        "Finally, we show that SteamCharts data, combined with Steam's public APIs, can support rigorous causal inference "
        "for a wide range of research questions in game markets. The data collection pipeline and quality verification "
        "procedures we document can be adapted to study other interventions such as sales, content updates, or platform changes."
    )
    add_normal_text(doc, text)
    
    add_heading(doc, '4.4 Limitations and Future Research', level=2)
    
    text = (
        "Several limitations should inform interpretation of our findings. First, our outcome measure—average concurrent "
        "players—captures only one dimension of player engagement. Future research should examine other metrics such as "
        "total playtime, session frequency, retention rates, and revenue. Patches may affect these outcomes differently, "
        "and the business case for patching depends on revenue impacts that we do not observe."
    )
    add_normal_text(doc, text)
    
    text = (
        "Second, we treat all 'major patches' as homogeneous interventions, but patch content varies enormously—from "
        "pure bug fixes to expansive new content releases. Future work should distinguish patch types and estimate "
        "content-specific effects. We expect that patches adding substantive new features or content yield larger "
        "engagement effects than technical updates."
    )
    add_normal_text(doc, text)
    
    text = (
        "Third, our sample selection process—requiring complete SteamCharts data for all time periods—likely excludes "
        "smaller games with minimal player tracking. Results may not generalize to indie games with player counts below "
        "SteamCharts' tracking threshold. Moreover, our focus on Steam PC games means findings may not apply to mobile, "
        "console, or other gaming platforms with different player dynamics."
    )
    add_normal_text(doc, text)
    
    text = (
        "Fourth, our observation window is limited to 5 months maximum. We cannot assess whether the measured effects "
        "persist over longer horizons or dissipate as players exhaust new content. Longitudinal studies tracking games "
        "for 6-12 months post-patch would clarify the durability of treatment effects."
    )
    add_normal_text(doc, text)
    
    text = (
        "Future research directions include: (1) heterogeneous effects by game genre, distinguishing multiplayer "
        "competitive games from single-player narrative games; (2) mechanisms underlying patch effects, separating "
        "actual content improvement from media attention and community hype; (3) interaction effects between patches "
        "and other marketing activities such as sales or promotions; (4) spillover effects on competing games in the "
        "same genre; and (5) optimal patching frequency and timing strategies."
    )
    add_normal_text(doc, text)
    
    add_heading(doc, '4.5 Concluding Remarks', level=2)
    
    text = (
        "This study provides credible causal evidence that major game patches modestly but significantly increase player "
        "engagement in Steam games. The 6% effect we estimate is policy-relevant for developers making post-launch support "
        "decisions, while highlighting that patches alone cannot compensate for fundamental game quality deficits. "
        "Methodologically, we demonstrate the severe selection bias present in naive comparisons and the necessity of "
        "fixed effects approaches for credible causal inference in observational game data."
    )
    add_normal_text(doc, text)
    
    text = (
        "As the gaming industry increasingly adopts 'games as a service' models with continuous post-launch updates, "
        "understanding the causal effects of these interventions becomes crucial for strategic planning. Our findings "
        "suggest that while patches contribute to player retention, their effects are incremental rather than transformative. "
        "Developer resources may be better invested in maximizing initial quality than in reactive patching, though the "
        "optimal balance depends on game-specific characteristics and business models that merit further investigation."
    )
    add_normal_text(doc, text)
    
    # ========================================
    # SAVE DOCUMENT
    # ========================================
    
    output_path = 'Steam_Patches_DiD_Analysis_Paper.docx'
    doc.save(output_path)
    print(f"\n{'='*80}")
    print(f"✓ Academic paper generated successfully!")
    print(f"{'='*80}")
    print(f"\nOutput file: {output_path}")
    print(f"Location: {os.path.abspath(output_path)}")
    print(f"\n{'='*80}")
    print("Document Structure:")
    print("  - Abstract")
    print("  - Chapter 1: Data Sources and Sample Construction")
    print("  - Chapter 2: Difference-in-Differences Methodology")
    print("  - Chapter 3: Empirical Results")
    print("  - Chapter 4: Conclusion")
    print(f"{'='*80}\n")
    
    return output_path

if __name__ == '__main__':
    create_academic_paper()
